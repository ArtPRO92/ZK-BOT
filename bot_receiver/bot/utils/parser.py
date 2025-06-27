import io
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import openpyxl

# Указание пути к Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """Предобработка изображения для лучшего OCR"""
    image = image.convert("L")  # grayscale
    image = image.point(lambda x: 0 if x < 180 else 255, mode='1')  # бинаризация
    image = image.resize((image.width * 2, image.height * 2))  # увеличение
    return image

def extract_text_from_docx(content: bytes) -> str:
    try:
        file_stream = io.BytesIO(content)
        doc = Document(file_stream)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_texts.append(" | ".join(row_text))

        all_text = paragraphs + table_texts
        print(f"🔍 Параграфов: {len(paragraphs)}, строк из таблиц: {len(table_texts)}")
        return "\n".join(all_text)

    except Exception as e:
        print("❌ Ошибка при разборе .docx:", e)
        return ""

def extract_text_from_pdf(content: bytes) -> str:
    text = ""
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                text += page_text + "\n"
            else:
                print(f"📄 Страница {page_num} не содержит текста — запускаем OCR")
                rotation = page.rotation or 0

                if rotation:
                    mat = fitz.Matrix(1, 1).preRotate(-rotation)
                    pix = page.get_pixmap(dpi=300, matrix=mat)
                else:
                    pix = page.get_pixmap(dpi=300)

                img = Image.open(io.BytesIO(pix.tobytes("png")))
                img = preprocess_image_for_ocr(img)
                ocr_text = pytesseract.image_to_string(img, lang="rus+eng")
                text += ocr_text + "\n"
    return text

def extract_text_from_image(content: bytes) -> str:
    image = Image.open(io.BytesIO(content))
    image = preprocess_image_for_ocr(image)
    return pytesseract.image_to_string(image, lang="eng+rus")

def extract_text_from_xlsx(content: bytes) -> str:
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        text_rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell is not None]
                if values:
                    text_rows.append(" | ".join(values))
        print(f"📊 Excel строк: {len(text_rows)}")
        return "\n".join(text_rows)
    except Exception as e:
        print("❌ Ошибка при разборе .xlsx:", e)
        return ""
